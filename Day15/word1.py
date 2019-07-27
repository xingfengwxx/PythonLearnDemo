#!/usr/bin/env python 3.7
# -*- coding: utf-8 -*-
# @Time       : 2019/7/27 10:06
# @Author     : wangxingxing
# @Email      : xingfengwxx@gmail.com 
# @File       : word1.py
# @Software   : PyCharm
# @Description: 读取Word文件

from docx import Document

doc = Document('./res/用函数还是用复杂的表达式.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)
print(''.join(content))
